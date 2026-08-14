"""Generate 'Aviation_NLP_Pipeline_Report.docx' - project summary + full implementation code.

The report embeds the current PROJECT_OVERVIEW.txt, the latest run metrics
(accuracy, CV F1, best hyperparameters) and the complete source of every file
in the production-grade `safety_nlp_pipeline/` module.
"""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Aviation_NLP_Pipeline_Report.docx"
PIPE = ROOT / "safety_nlp_pipeline"
OVERVIEW = ROOT / "PROJECT_OVERVIEW.txt"

CODE_FILES = [
    "safety_nlp_pipeline/main.py",
    "safety_nlp_pipeline/config.py",
    "safety_nlp_pipeline/app_streamlit.py",
    "safety_nlp_pipeline/src/__init__.py",
    "safety_nlp_pipeline/src/data_fetcher.py",
    "safety_nlp_pipeline/src/preprocessor.py",
    "safety_nlp_pipeline/src/trainer.py",
    "safety_nlp_pipeline/src/evaluator.py",
    "safety_nlp_pipeline/src/rag_explainer.py",
    "safety_nlp_pipeline/src/analyst.py",
    "safety_nlp_pipeline/src/monitor.py",
    "safety_nlp_pipeline/src/report_generator.py",
    "safety_nlp_pipeline/requirements.txt",
    "safety_nlp_pipeline/.streamlit/config.toml",
]

ACCENT = RGBColor(0x0F, 0x6B, 0x9E)
CODE_BG = "F2F4F6"
CODE_FONT = "Consolas"


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc, code: str):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x33)
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        pf.left_indent = Inches(0.08)
        shade(p, CODE_BG)
    # spacing after the block
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _read_json(rel_path: str) -> dict:
    """Best-effort JSON read; returns {} on any failure."""
    try:
        return json.loads((ROOT / rel_path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def load_latest_metrics() -> dict:
    """Pull the most recent training/evaluation numbers from data/."""
    train = _read_json("safety_nlp_pipeline/data/training_config.json")
    met = _read_json("safety_nlp_pipeline/data/metrics.json")
    return {
        "accuracy": met.get("accuracy"),
        "n_classes": met.get("n_classes"),
        "test_size": met.get("test_size"),
        "best_cv_f1": train.get("best_cv_f1_weighted"),
        "best_params": train.get("best_params", {}),
        "n_train": train.get("n_train"),
    }


def add_latest_results(doc, metrics: dict) -> None:
    doc.add_heading("1.4 Latest run results", level=2)
    present = any(v is not None for v in metrics.values())
    if not present:
        doc.add_paragraph("No persisted metrics found yet - run `python main.py` "
                          "inside safety_nlp_pipeline/ to generate them.")
        return
    lines = []
    if metrics.get("accuracy") is not None:
        lines.append(f"Test accuracy            : {metrics['accuracy']:.1%}")
    if metrics.get("best_cv_f1") is not None:
        lines.append(f"Best CV F1 (weighted)    : {metrics['best_cv_f1']:.3f}")
    if metrics.get("n_classes") is not None:
        lines.append(f"Classes                  : {metrics['n_classes']}")
    if metrics.get("n_train") is not None:
        lines.append(f"Training rows            : {metrics['n_train']}")
    if metrics.get("test_size") is not None:
        lines.append(f"Test rows                : {metrics['test_size']}")
    bp = metrics.get("best_params") or {}
    if bp:
        lines.append("Best hyperparameters     : "
                     + "; ".join(f"{k} = {v}" for k, v in bp.items()))
    add_code_block(doc, "\n".join(lines))


def main():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---------- Title ----------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Aviation & Power-Grid Safety NLP Pipeline")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Production-grade headless pipeline — NASA ASRS · NERC power grid · "
                  "TF-IDF · SGD · GridSearchCV · RAG · HTML report · Streamlit dashboard · "
                  "real-time monitoring & alerting")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run("Project Summary, Updated Project Overview & Full Implementation Code (DOCX)")
    r.font.size = Pt(11)
    r.font.italic = True

    doc.add_paragraph()

    # ---------- 1. Project summary ----------
    doc.add_heading("1. Project Summary", level=1)
    doc.add_paragraph(
        "This project is a production-grade, HEADLESS NLP pipeline that runs end-to-end "
        "with zero user intervention (`python main.py`). It classifies real safety-incident "
        "reports from two public domains - NASA ASRS aviation incident reports and NERC "
        "power-grid event analysis reports - and produces a comprehensive, self-contained "
        "HTML report (model metrics, confusion matrix, RAG evidence, data quality insights). "
        "A Streamlit web dashboard with a warm, creamy-light theme makes the results "
        "accessible to stakeholders who are not comfortable with the command line."
    )

    doc.add_heading("1.1 Key features", level=2)
    for item in [
        "One command, no intervention: python main.py runs fetch -> preprocess -> train -> evaluate -> RAG -> HTML report.",
        "Model robustness: StratifiedKFold cross-validation + GridSearchCV hyperparameter tuning (alpha x class_weight) over TF-IDF (bigrams) + SGD(log-loss).",
        "Fault tolerance: each data source is fetched independently with graceful per-domain fallback to cached data; the pipeline is idempotent (skips fetch if a cached CSV exists).",
        "Live data collection: NASA ASRS reports (Hugging Face datasets-server) + 12 NERC event-analysis PDFs (pypdf + NLTK sentence tokenization), merged into one domain-tagged dataset.",
        "Rich HTML report: Jinja2 template with dataset statistics, data-quality audit, best hyperparameters, per-class metrics table, embedded confusion-matrix + class distribution plots, and RAG evidence examples with similarity bars.",
        "Batch RAG explainability: a sample of the test set is explained with the top-3 most similar historical reports per prediction - an auditable, explainable-by-example system.",
        "Streamlit web dashboard (creamy light theme): Overview, Model Performance, RAG Explorer (live predictions + evidence), Data Assistant and Live Alerts tabs.",
        "Real-time incident monitoring & alerting (src/monitor.py): ingests new reports as they arrive, classifies on-the-fly, scores risk (critical/high/medium) and raises alerts with RAG evidence - from a watched folder, appended dataset rows, the live NTSB aviation feed and the UK Power Networks live-faults feed.",
        "Alert de-duplication survives restarts (data/monitor_state.json), and the Live Alerts dashboard tab color-codes critical/high rows and shows RAG evidence per alert.",
        "Keyless data assistant (no LLM/API key): data-quality issues, class balance, domain split, safety-criticality breakdown and risk-phrase scanning with pandas.",
        "Logging: every step writes to the console and pipeline.log for full traceability.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.2 Pipeline architecture", level=2)
    arch = (
        "data/real_safety_dataset.csv   (NASA ASRS + NERC, ~3,700 domain-tagged reports)\n"
        "        |\n"
        "  [1/6] FETCH   (live: ASRS datasets-server + NERC PDFs -> clean -> merge by domain)\n"
        "        |\n"
        "  [2/6] PREPROCESS  (NLTK tokenize -> stopword removal -> lemmatize, per-document)\n"
        "        |\n"
        "  [3/6] TRAIN   (TF-IDF bigrams, max_features=5000 -> SGD/log-loss,\n"
        "        |        GridSearchCV: alpha x class_weight, StratifiedKFold CV,\n"
        "        |        best model + vectorizer saved to data/*.pkl)\n"
        "  [4/6] EVALUATE (classification report, confusion-matrix heatmap,\n"
        "        |         class-distribution plot -> data/plots/*.png, persisted metrics)\n"
        "  [5/6] RAG EXPLAINABILITY (test-set sample -> predict each report ->\n"
        "        |                     batch cosine similarity vs all training\n"
        "        |                     narratives -> top-3 evidence spans)\n"
        "  [6/6] HTML REPORT (Jinja2 self-contained page -> reports/pipeline_report.html)\n"
        "        |\n"
        "  [Web] STREAMLIT DASHBOARD (app_streamlit.py, creamy light theme)\n"
        "        |\n"
        "  [MON] REAL-TIME MONITOR & ALERTER (src/monitor.py, python -m src.monitor)\n"
        "        |   watches new_incidents/ + appended dataset rows + live NTSB / UKPN feeds\n"
        "        |   -> classify on-the-fly -> risk score -> alert with RAG evidence\n"
        "        `-> data/alerts.csv -> Streamlit 'Live Alerts' tab"
    )
    add_code_block(doc, arch)

    add_latest_results(doc, load_latest_metrics())

    doc.add_heading("1.5 How to run", level=2)
    add_code_block(
        doc,
        "cd safety_nlp_pipeline\n"
        "pip install -r requirements.txt\n"
        "python main.py                        # full pipeline, one command\n"
        "# open reports/pipeline_report.html in a browser\n\n"
        "streamlit run app_streamlit.py        # web dashboard (http://localhost:8501)\n\n"
        "# real-time monitoring & alerting:\n"
        "python -m src.monitor                 # continuous monitor (needs trained model)\n"
        "python -m src.monitor --once --no-api # single scan, no live feeds\n"
        "python main.py --monitor --poll 30    # train, then start monitoring\n\n"
        "# optional pipeline flags:\n"
        "python main.py --force-refresh        # re-download data\n"
        "python main.py --no-fetch             # use cached CSV only\n"
        "python main.py --no-rag               # skip RAG explainability\n"
        "python main.py --samples 200          # RAG test samples",
    )

    # ---------- 2. Updated project overview ----------
    doc.add_heading("2. Updated Project Overview", level=1)
    doc.add_paragraph(
        "The full, current project overview is reproduced below (from "
        "PROJECT_OVERVIEW.txt) and reflects everything implemented in the "
        "restructured production build."
    )
    overview_text = OVERVIEW.read_text(encoding="utf-8") if OVERVIEW.exists() else ""
    add_code_block(doc, overview_text)

    # ---------- 3. Project structure ----------
    doc.add_heading("3. Project Structure", level=1)
    add_code_block(
        doc,
        "Aviation_NLP_Project/\n"
        "|-- make_report.py             builds this DOCX report\n"
        "|-- PROJECT_OVERVIEW.txt       full updated project overview\n"
        "|-- app.py, pipeline/          legacy Textual TUI (kept as a developer tool)\n"
        "`-- safety_nlp_pipeline/\n"
        "    |-- README.md\n"
        "    |-- requirements.txt\n"
        "    |-- config.py              all parameters (paths, model settings, ...)\n"
        "    |-- main.py                single entry point - runs the full pipeline\n"
        "    |-- app_streamlit.py       Streamlit web dashboard (creamy light theme)\n"
        "    |-- .streamlit/config.toml dashboard theme configuration\n"
        "    |-- src/\n"
        "    |   |-- data_fetcher.py    downloads ASRS (HF) + NERC (PDFs) -> CSV\n"
        "    |   |-- preprocessor.py    NLTK tokenization, stopwords, lemmatization\n"
        "    |   |-- trainer.py         TF-IDF + SGD classifier (log-loss) with GridSearchCV\n"
        "    |   |-- evaluator.py       classification report, confusion matrix, plots\n"
        "    |   |-- rag_explainer.py   batch + single-query semantic retrieval (cosine)\n"
        "    |   |-- analyst.py         keyless data quality & safety analysis\n"
        "    |   |-- monitor.py         real-time incident monitoring & alerting\n"
        "    |   `-- report_generator.py self-contained HTML report (Jinja2)\n"
        "    |-- new_incidents/         drop-in folder for new CSV/TXT reports\n"
        "    |-- data/                  auto-created; CSV, models, plots, metrics, alerts\n"
        "    `-- reports/               pipeline_report.html",
    )

    # ---------- 4. Implementation code ----------
    doc.add_heading("4. Implementation Code", level=1)
    doc.add_paragraph(
        "Complete source of every file in the production pipeline. Each code block is "
        "rendered in monospace (Consolas) with a shaded background for readability."
    )
    for rel in CODE_FILES:
        path = ROOT / rel
        if not path.exists():
            doc.add_heading(f"4.{CODE_FILES.index(rel) + 1}  {rel}  (MISSING)",
                            level=2)
            continue
        doc.add_heading(f"4.{CODE_FILES.index(rel) + 1}  {rel}", level=2)
        add_code_block(doc, path.read_text(encoding="utf-8"))

    # ---------- 5. Dependencies ----------
    doc.add_heading("5. Dependencies", level=1)
    add_code_block(
        doc,
        "pandas>=2.2         Dataframes for dataset processing & assistant analysis\n"
        "scikit-learn>=1.5   TF-IDF, SGDClassifier, GridSearchCV, cosine similarity\n"
        "datasets>=3.0       Hugging Face datasets-server client (aviation source)\n"
        "joblib>=1.4         Model & vectorizer pickling\n"
        "nltk>=3.8           Tokenization, stopwords, lemmatization, sentence split\n"
        "pypdf>=4.0          NERC PDF text extraction (power-grid source)\n"
        "requests>=2.31      HTTP downloads (HF datasets-server + NERC PDFs)\n"
        "matplotlib>=3.8     Confusion matrix & class distribution plots\n"
        "seaborn>=0.13       Heatmap rendering\n"
        "jinja2>=3.1         HTML report templates\n"
        "streamlit>=1.37     Web dashboard (creamy light theme)",
    )

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("— End of report —")
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x9A, 0xA0, 0xAA)

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
